import sys
import json


from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_UNDERLINE


USAGE = "USAGE: \npython3 json_to_docx.py <json_file> <with_thoughts>\nwhere\n* <json_file> is the filename of the json file of the run (without the extention)\n* <with_thoughts> is either True or False, and determines whether the Thoughts are processed or not."

# Define color mappings for thought types
color_mapping = {
    "watsonian": RGBColor(128, 0, 128),  # Purple
    "doylist": RGBColor(255, 0, 0),      # Red
    "meta": RGBColor(0, 128, 0),         # Green
    "comment": RGBColor(0, 0, 255)       # Blue
}

# Function to add underlined style
def apply_underlined_style(run):
    run.font.underline = WD_UNDERLINE.SINGLE

# Function to add colored and styled text
def add_colored_text(paragraph, text, color, italic=False, underline=False, mono=False):
    run = paragraph.add_run(text)
    run.font.color.rgb = color
    if italic:
        run.italic = True
    if underline:
        apply_underlined_style(run)
    if mono:
        run.font.name = 'Courier New'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

def process_thoughts(doc, L):
    doc.add_heading('Thoughts', level=3)
    if 'thoughts' in L:
        for thought_group in L['thoughts']:
            paragraph = doc.add_paragraph()
            paragraph.style = "List Bullet"
            for thought in thought_group:
                color = color_mapping.get(thought['type'], RGBColor(0, 0, 0))  # Default to black if type is unknown
                italic = thought['type'] == 'doylist'
                underline = thought.get('longterm', False)
                mono = thought['type'] == 'comment'
                add_colored_text(paragraph, thought['text']+" ", color, italic, underline, mono)
    else:
        paragraph = doc.add_paragraph()
        paragraph.add_run("<No Thoughts>")

def main (filename, with_thoughts):
    if with_thoughts == None:
        raise ValueError

    # Load the JSON file
    with open(filename+".json", 'r') as file:
        data = json.load(file)

    # Create a new Document
    doc = Document()
    doc.add_heading(data['title'], level=0)
    doc.add_page_break()

    # Process each step
    step_number = 1
    for step in data['steps']:
        
        # Add a page break before each new step, except the first one
        if step_number > 1:
            doc.add_page_break()

        # Add step heading
        doc.add_heading(f'Step {step_number}', level=2)
        step_number += 1
        
        # Process thoughts
        if with_thoughts:
            process_thoughts(doc, step)

        # Process prompt
        doc.add_heading('Prompt', level=3)
        if 'prompt' in step:        
            doc.add_paragraph(step['prompt']['text'])
            if with_thoughts:
                process_thoughts(doc, step['prompt'])
        else:
            doc.add_paragraph("<Skipped Prompt>")
        
        # Process action
        if 'action' in step:
            doc.add_heading('Action', level=3)
            doc.add_paragraph(step['action']['text'])
            
            if with_thoughts:
                process_thoughts(doc, step['action'])

            doc.add_heading('Outcome', level=3)
            if 'outcome' in step['action']:
                doc.add_paragraph(step['action']['outcome'])
            else:
                doc.add_paragraph("<Skipped Outcome>")

    # Save the document
    if with_thoughts:
        new_filename = filename + "_with_thoughts.docx"
    else:
        new_filename = filename + ".docx"
    doc.save(new_filename)

try:
    with_thoughts = None
    if sys.argv[2]=="True":
        with_thoughts = True
    if sys.argv[2]=="False":
        with_thoughts = False
    main(sys.argv[1], with_thoughts)
except (IndexError, ValueError):
    print(USAGE)
